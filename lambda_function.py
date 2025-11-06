import os
import io
import json
from typing import Optional
import boto3
import pdfplumber
import docx
import mimetypes
import time
import logging
import re
from datetime import datetime
from decimal import Decimal
from botocore.exceptions import ClientError
from boto3.dynamodb.conditions import Attr
import urllib.parse
from functools import lru_cache
from boto3 import Session

# Configure structured logging
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# Initialize AWS clients
dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
s3_client = boto3.client('s3')

# OpenAI API Key - fetch from environment variables
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY')

if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY environment variable is required but not set")

# Load documentation files with caching
@lru_cache(maxsize=1)
def load_documentation():
    """Load documentation files with caching"""
    docs = {}
    try:
        with open('model_process.md', 'r') as file:
            docs['model_process'] = file.read()
        print("Successfully loaded model_process.md file")

        with open('level_finding.md', 'r') as file:
            docs['level_finding'] = file.read()
        print("Successfully loaded level_finding.md file")
        
        # Remove category.md as it's not needed for document analysis
        print("Skipping category.md - using document-based risk assessment")
    except Exception as e:
        print(f"Error loading documentation files: {e}")
        docs['model_process'] = "Model process documentation not available"
        docs['level_finding'] = "Level finding documentation not available"

    return docs

# Get documentation content
docs = load_documentation()
level_finding_content = docs['level_finding']
model_process_content = docs['model_process']

# Create document-based risk assessment guidance
document_risk_guidance = """
DOCUMENT-BASED RISK ASSESSMENT CRITERIA:

LOW-RISK AI:
- Basic automation or efficiency improvements
- Limited data processing scope
- No personal/sensitive data involved
- Simple business process enhancement
- Minimal regulatory implications

MEDIUM-RISK AI:
- Moderate business process automation
- Some personal data processing
- Industry-specific compliance considerations
- Moderate technical complexity
- Potential for operational impact

HIGH-RISK AI:
- Critical business decision automation
- Extensive personal/sensitive data processing
- High regulatory compliance requirements (GDPR, HIPAA, financial regulations)
- Safety-critical applications
- Significant operational or financial impact
- Public-facing AI systems
- AI systems affecting individual rights or opportunities

PROHIBITED AI:
- Biometric identification in public spaces
- AI systems for social scoring
- Subliminal techniques to manipulate behavior
- Exploitation of vulnerabilities (age, disability)
- Real-time emotion recognition in workplace/education
"""

def convert_floats_to_decimal_for_dynamodb(obj):
    """Convert float values to Decimal for DynamoDB compatibility"""
    if isinstance(obj, float):
        return Decimal(str(obj))
    elif isinstance(obj, dict):
        return {k: convert_floats_to_decimal_for_dynamodb(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [convert_floats_to_decimal_for_dynamodb(item) for item in obj]
    else:
        return obj

def clean_model_process(model_process):
    """
    Clean the model process to remove any markdown formatting and ensure proper structure
    """
    # Remove markdown headers (# symbols)
    model_process = re.sub(r'^#+\s*', '', model_process, flags=re.MULTILINE)

    # Remove bullet points
    model_process = re.sub(r'^\s*[-•*]\s*', '', model_process, flags=re.MULTILINE)

    # Remove bold markdown
    model_process = re.sub(r'\*\*(.*?)\*\*', r'\1', model_process)

    # Split by common stage names to rebuild properly
    stage_names = [
        # Machine Learning stages
        "Data Collection", "Data Preprocessing", "Feature Engineering",
        "Model Training", "Dataset Preparation", "Model Architecture",
        "Hyperparameter Optimization", "Validation", "Model Evaluation",
        "Deployment", "Monitoring", "Feedback Loop", "Retraining", 
        "Audit Trail", "Compliance Check", "Maintenance",
        # Gen AI stages (updated universal names)
        "Requirements Analysis", "Data Preparation", "Content Processing",
        "Foundation Setup", "Model Configuration", "Prompt Engineering",
        "Implementation & Deployment", "Performance Monitoring", "Governance & Compliance",
        # Legacy GenAI stage names (for backward compatibility)
        "Requirements Gathering", "Data Ingestion", "Data Chunking & Processing",
        "Knowledge Base Setup", "Model Selection & Configuration", "Prompt Design & Testing"
    ]

    # Create a pattern to match stage names with various formatting
    stages = {}
    for stage_name in stage_names:
        # Find stage content using regex
        pattern = rf'{stage_name}[:\s]*([^(Data |Feature |Model |Dataset |Hyperparameter |Validation|Deployment|Monitoring|Feedback |Retraining|Audit |Compliance|Maintenance|Content |Prompt )]+)'
        match = re.search(pattern, model_process, re.IGNORECASE | re.DOTALL)
        if match:
            content = match.group(1).strip()
            # Clean up the content
            content = re.sub(r'\s+', ' ', content)  # Replace multiple spaces with single space
            content = content.replace('\\n', ' ').replace('\n', ' ')  # Remove line breaks within content
            # Remove any remaining formatting characters
            content = re.sub(r'[•\-\*]', '', content).strip()

            if content and len(content) > 20:  # Only add if there's substantial content
                stages[stage_name] = f"{stage_name}: {content}"

    # Rebuild the model process with proper formatting
    if stages:
        # Use only the stages that were found
        rebuilt_process = []
        for stage_name in stage_names:
            if stage_name in stages:
                rebuilt_process.append(stages[stage_name])

        # Join with single newline (not double)
        return '\n'.join(rebuilt_process)

    # If parsing failed, return original with basic cleanup
    return model_process.replace('\\n\\n', '\n').replace('\\n', '\n').strip()

def preprocess_usecase_data(usecase_data):
    """Preprocess usecase data to ensure proper formatting before storage"""
    if 'modelProcess' in usecase_data and usecase_data['modelProcess']:
        model_process = usecase_data['modelProcess']

        # Check if already has proper single newline formatting
        if "\n" in model_process and "\n\n" not in model_process:
            print("Model process already has proper formatting")
            return usecase_data

        # If double newlines exist, convert to single
        if "\n\n" in model_process:
            usecase_data['modelProcess'] = model_process.replace("\n\n", "\n")
            print("Converted double newlines to single newlines")
            return usecase_data

        # Otherwise, try to add newlines between stages
        stage_prefixes = [
            "Data Collection:", "Data Preprocessing:", "Feature Engineering:",
            "Model Training:", "Dataset Preparation:", "Model Architecture:",
            "Hyperparameter Optimization:", "Validation:", "Model Evaluation:",
            "Deployment:", "Monitoring:", "Feedback Loop:", "Retraining:",
            "Audit Trail:", "Compliance Check:", "Maintenance:",
            # Updated Gen AI stage prefixes
            "Requirements Analysis:", "Data Preparation:", "Content Processing:",
            "Foundation Setup:", "Model Configuration:", "Prompt Engineering:",
            "Implementation & Deployment:", "Performance Monitoring:", "Governance & Compliance:",
            # Legacy stage prefixes for backward compatibility
            "Requirements Gathering:", "Data Ingestion:", "Data Chunking & Processing:",
            "Knowledge Base Setup:", "Model Selection & Configuration:", "Prompt Design & Testing:"
        ]

        processed_model_process = model_process
        for i, prefix in enumerate(stage_prefixes):
            if i == 0:  # Don't add newline before the first stage
                continue

            if prefix in processed_model_process:
                # Add single newline for separation
                processed_model_process = processed_model_process.replace(prefix, f"\n{prefix}")
                print(f"Added newline before stage: {prefix}")

        usecase_data['modelProcess'] = processed_model_process
        print("Successfully preprocessed model process with proper formatting")

    return usecase_data

def determine_ai_category_and_approach_with_llm(document_text):
    """Determine AI category and technical approach based on document content"""
    print("Determining AI category and approach using LLM analysis")

    prompt = f"""
Analyze the following COMPLETE document content and identify BOTH:
1. AI Category → One of:
   - Machine Learning
   - AI Workflow Agents
   - Agentic AI
2. Approach → Based on tech stack and complexity:
   - "Next.js + ADK" (simple use case)
   - "Next.js + ADK + MCP" (MCP multi-agent use case)
   - "Custom AI Stack" (if neither matches)

COMPLETE DOCUMENT CONTENT:
{document_text}

CRITERIA:
- Category logic:
  - Machine Learning: Prediction, classification, regression, forecasting, optimization
  - AI Workflow Agents: Step-based or workflow-driven tasks, chatbots, RAG, summarization, single/few-shot flows
  - Agentic AI: Multi-agent, orchestration, decision-making, reasoning, goal-directed autonomy
- Approach logic:
  - Next.js + ADK → If mentions frontend AI app, ADK, basic integration, no orchestration
  - Next.js + ADK + MCP → If mentions MCP, multi-agent, orchestration, workflows, agents collaborating
  - Custom AI Stack → Any other tech stack (e.g., Python FastAPI, LangChain, Spring Boot etc.)

INSTRUCTIONS:
1. Analyze ENTIRE content for both category and approach.
2. Respond strictly in the format:

Category: [Category]
Approach: [Approach]
Reason: [1-2 sentence explanation]
"""

    try:
        system_prompt = "You are an AI categorization and approach expert. Classify both AI category and technical approach from the document."

        response = call_llm(prompt, system_prompt, use_json_format=False, max_tokens=300)

        # Default fallback values
        category, approach = "AI Workflow Agents", "Custom AI Stack"

        # --- Category handling ---
        if "Machine Learning" in response:
            category = "Machine Learning"
        elif "Agentic AI" in response:
            category = "Agentic AI"
        elif "AI Workflow Agents" in response:
            category = "AI Workflow Agents"

        # --- Approach handling ---
        if "Next.js + ADK + MCP" in response:
            approach = "MCP-Oriented Orchestration"
        elif "Next.js + ADK" in response:
            approach = "Next.js with ADK"
        elif "Custom AI Stack" in response:
            approach = "Custom AI Stack"

        return {"category": category, "approach": approach, "raw_response": response}

    except Exception as e:
        print(f"Error in LLM determination: {e}")

        # ✅ Call fallback keyword-based categorization
        category = determine_ai_category_based_on_data_fallback(document_text)

        # ✅ Simple keyword-based approach detection
        text_lower = document_text.lower()
        if "mcp" in text_lower or "multi-agent" in text_lower or "orchestration" in text_lower:
            approach = "MCP-Oriented Orchestration"
        elif "next.js" in text_lower and "adk" in text_lower:
            approach = "Next.js with ADK"
        else:
            approach = "Custom AI Stack"

        return {"category": category, "approach": approach, "raw_response": "Fallback logic applied"}

def call_llm(prompt: str, system_prompt: Optional[str] = None, model_id: str = "gpt-4o", use_json_format: bool = True, max_tokens: int = 4000) -> str:
    """Call OpenAI GPT-4o via OpenAI API with configurable max_tokens"""
    import requests

    url = "https://api.openai.com/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {OPENAI_API_KEY}"
    }

    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.append({"role": "user", "content": prompt})

    data = {
        "model": model_id,
        "messages": messages,
        "temperature": 0.3 if not use_json_format else 0.5,
        "max_tokens": max_tokens
    }

    if use_json_format:
        data["response_format"] = {"type": "json_object"}

    try:
        start_time = time.time()
        response = requests.post(url, headers=headers, json=data)
        response.raise_for_status()
        result = response.json()

        # Log latency metrics
        latency = (time.time() - start_time) * 1000  # milliseconds
        print(f"OpenAI API call completed in {latency:.2f}ms")

        return result["choices"][0]["message"]["content"]
    except Exception as error:
        print(f"Error calling OpenAI API: {error}")
        raise Exception(f"LLM service error: {error}")

def determine_ai_category_based_on_data_fallback(document_text):
    """Fallback keyword-based categorization with Machine Learning, AI Workflow Agents, Agentic AI"""
    print("Using fallback keyword-based categorization")

    try:
        text_lower = document_text.lower()

        # ML keywords  
        ml_keywords = [
            'predict', 'prediction', 'classify', 'classification', 'cluster',
            'regression', 'forecast', 'anomaly detection', 'pattern recognition',
            'supervised learning', 'unsupervised learning', 'recommendation',
            'feature engineering', 'model training', 'algorithm',
            'neural network', 'deep learning', 'random forest', 'svm',
            'decision tree', 'ensemble', 'xgboost', 'logistic regression',
            'clustering', 'optimization', 'data mining'
        ]

        # AI Workflow Agent keywords (includes Gen AI patterns)
        workflow_keywords = [
            'chatbot', 'rag', 'retrieval augmented generation',
            'summarization', 'summarize', 'dialogue', 'conversation',
            'workflow', 'approval process', 'single shot', 'few shot',
            'content creation', 'form processing', 'automation step'
        ]

        # Agentic AI keywords
        agentic_keywords = [
            'agent', 'autonomous', 'multi-agent', 'collaboration',
            'orchestration', 'decision-making', 'reasoning', 'goal-oriented',
            'workflow automation', 'adaptive planning',
            'self-improving', 'tool use', 'api orchestration', 'planner',
            'mcp', 'model context protocol'
        ]

        ml_score = sum(1 for keyword in ml_keywords if keyword in text_lower)
        workflow_score = sum(1 for keyword in workflow_keywords if keyword in text_lower)
        agentic_score = sum(1 for keyword in agentic_keywords if keyword in text_lower)

        print(f"Keyword analysis - ML: {ml_score}, Workflow: {workflow_score}, Agentic: {agentic_score}")

        if ml_score > workflow_score and ml_score > agentic_score:
            return "Machine Learning"
        elif workflow_score > ml_score and workflow_score > agentic_score:
            return "AI Workflow Agents"
        elif agentic_score > ml_score and agentic_score > workflow_score:
            return "Agentic AI"
        else:
            # ✅ Fallback = Workflow Agents (base/default)
            return "AI Workflow Agents"

    except Exception as e:
        print(f"Error in fallback categorization: {e}")
        return "AI Workflow Agents"  # Safe default
    

def fetch_all_dynamodb_records(table_name, dynamodb_resource):
        """Fetch all records from a DynamoDB table, handling pagination."""
        try:
            table = dynamodb_resource.Table(table_name)
            all_records = []
            scan_kwargs = {}
            while True:
                response = table.scan(**scan_kwargs)
                items = response.get('Items', [])
                all_records.extend(items)
                if 'LastEvaluatedKey' in response:
                    scan_kwargs['ExclusiveStartKey'] = response['LastEvaluatedKey']
                else:
                    break
            print(f"Fetched {len(all_records)} records from {table_name}")
            return all_records
        except Exception as e:
            print(f"Error fetching records from {table_name}: {e}")
            return []

# def extract_app_and_stage_from_event(event):
#         """Extract app_name, stage_name, and s3_bucket_arn from the S3 event structure."""
#         try:
#             s3_bucket_arn = event['Records'][0]['s3']['bucket']['arn']
#             print(f"Extracted bucket ARN: {s3_bucket_arn}")
#         except (KeyError, IndexError) as e:
#             raise ValueError(f"Failed to extract S3 information from event: {e}")

#         bucket_name = s3_bucket_arn.split(':')[-1]
#         bucket_parts = bucket_name.split('-')
#         if len(bucket_parts) < 2:
#             raise ValueError(f"Bucket name '{bucket_name}' does not follow expected format")
#         app_name = bucket_parts[0]
#         stage_name = bucket_parts[1]
#         print(f"Derived app_name={app_name}, stage_name={stage_name}")
#         return app_name, stage_name, s3_bucket_arn

def extract_app_stage_tenant_from_bucket(bucket_name: str):
    """
    Extracts app_name, stage_name, and tenant_id_suffix from the S3 bucket name.
    Expected format: <app>-<stage>-<last6OfTenantId>
    Example: fusefy-staging-123abc → app_name=fusefy, stage_name=staging, tenant_id_suffix=123abc
    """
    print(f"📦 Deriving app/stage/tenant from bucket: {bucket_name}")
    parts = bucket_name.split('-')

    if len(parts) < 3:
        raise ValueError(f"Bucket name '{bucket_name}' does not follow expected format")

    app_name = parts[0]
    stage_name = parts[1]
    tenant_id_suffix = parts[-1]

    print(f"✅ Extracted app_name={app_name}, stage_name={stage_name}, tenant_id_suffix={tenant_id_suffix}")
    return app_name, stage_name, tenant_id_suffix


def summarize_document_with_llm(document_text):
        """
        Summarize the provided document text using OpenAI LLM, returning a documentSummary with key points.
        The summary should capture the main business use case and highlight key points relevant to the document's intent.
        """
        prompt = f"""
    You are an expert business analyst. Read the following business use case document and provide a concise summary (3-5 sentences) that captures the main purpose, business context, and technical approach. Then, list 3-5 key points or highlights from the document as bullet points.

    DOCUMENT:
    {document_text}

    RESPONSE FORMAT:
    Summary: <Concise summary here>
    Key Points:
    - <Key point 1>
    - <Key point 2>
    - <Key point 3>
    (Add more if relevant)
    """
        system_prompt = "You are a business analyst summarizing business use case documents for executive review."
        # Use higher max_tokens for richer summaries
        response = call_llm(prompt, system_prompt, use_json_format=False, max_tokens=600)
        # Optionally, you can parse the response to extract summary and key points if needed
        return {"documentSummary": response.strip()}



# def get_usecase_assessment_table_and_framework_id(event: dict, dynamodb) -> tuple[str, str | None]:
#     """
#     Extracts tenant-specific usecase assessment table name and framework ID from S3 event.

#     Returns:
#         (usecase_assessment_table, framework_id)
#     """
#     app_name, stage_name, s3_bucket_arn = extract_app_and_stage_from_event(event)

   
#     # Framework Table
#     framework_table = f"{stage_name}-{app_name}-frameworks"
#     try:
#         print(f"Looking up framework table: {framework_table}")
#         framework = dynamodb.Table(framework_table)
#         response = framework.scan(
#             FilterExpression=Attr('name').contains('NIST') & Attr('assessmentCategory').contains('AI Evaluation Engine'),
#             ProjectionExpression="id"
#         )
#         items = response.get('Items', [])
#         framework_id = items[0]['id'] if items else None
#         print(f"Framework ID found: {framework_id}")
#     except Exception as e:
#         print(f"Error scanning framework table {framework_table}: {e}")
#         framework_id = None  # fallback to None

#     # Tenant Details Table
#     tenant_details_table = f"{stage_name}-{app_name}-tenantDetails"
#     bucket_arn_last_6 = s3_bucket_arn[-6:]
#     tenant_id = None
#     try:
#         print(f"Looking up tenant details table: {tenant_details_table}, searching cloudId ending with {bucket_arn_last_6}")
#         tenant_table = dynamodb.Table(tenant_details_table)

#         response = tenant_table.scan(
#             FilterExpression=Attr('cloudId').contains(bucket_arn_last_6),
#             ProjectionExpression="cloudId"
#         )

#         for item in response.get('Items', []):
#             cloud_id = item.get('cloudId')
#             if cloud_id and cloud_id[-6:] == bucket_arn_last_6:
#                 tenant_id = cloud_id
#                 break

#         # Handle pagination if needed
#         while tenant_id is None and 'LastEvaluatedKey' in response:
#             response = tenant_table.scan(
#                 FilterExpression=Attr('cloudId').contains(bucket_arn_last_6),
#                 ProjectionExpression="cloudId",
#                 ExclusiveStartKey=response['LastEvaluatedKey']
#             )
#             for item in response.get('Items', []):
#                 cloud_id = item.get('cloudId')
#                 if cloud_id and cloud_id[-6:] == bucket_arn_last_6:
#                     tenant_id = cloud_id
#                     break

#         if not tenant_id:
#             raise ValueError(f"No matching tenant found for bucket ARN: {s3_bucket_arn}")

#         print(f"Resolved tenant_id={tenant_id}")

#     except Exception as e:
#         raise ValueError(f"Error scanning tenant details table {tenant_details_table}: {e}")

#     # Construct final table name
#     usecase_assessment_table = f"{stage_name}-{app_name}-usecaseAssessments-{tenant_id}"
#     print(f"Resolved usecase assessment table: {usecase_assessment_table}")

#     return usecase_assessment_table, framework_id


def get_usecase_assessment_table_and_framework_id(
    app_name: str, stage_name: str, tenant_id: str, dynamodb
) -> tuple[str, str | None]:
    """
    Returns tenant-specific usecase assessment table name and framework ID
    using app_name, stage_name, and tenant_id directly.
    """
    # Framework Table
    framework_table_name = f"{stage_name}-{app_name}-frameworks"
    try:
        framework_table = dynamodb.Table(framework_table_name)
        response = framework_table.scan(
            FilterExpression=Attr('name').contains('NIST') & Attr('assessmentCategory').contains('AI Evaluation Engine'),
            ProjectionExpression="id"
        )
        items = response.get('Items', [])
        framework_id = items[0]['id'] if items else None
        print(f"Framework ID found: {framework_id}")
    except Exception as e:
        print(f"Error scanning framework table {framework_table_name}: {e}")
        framework_id = None

    # Construct final usecase assessment table name
    usecase_table_name = f"{stage_name}-{app_name}-usecaseAssessments-{tenant_id}"
    print(f"Resolved usecase assessment table: {usecase_table_name}")

    return usecase_table_name, framework_id

def generate_design_document_prompt(document_text):
    """
    Create a comprehensive prompt for generating a design document
    
    Args:
        document_text (str): The content from the source document
        
    Returns:
        str: The formatted prompt for OpenAI
    """
    prompt = f"""
    MASTER AGENT PROMPT: Solution Architect and Project Planner
    AGENT ROLE: You are a Senior Solutions Architect and Agentic AI Specialist. You are an expert in Domain-Driven Design (DDD), Google Cloud Platform (GCP) architecture, and the NIST AI Risk Management Framework (AI RMF).

    GOAL: Analyze the provided Functional Requirement Document (FRD) and generate a comprehensive, structured technical design package. This package must adhere to all specified constraints, including strict NIST AI RMF compliance, focusing heavily on layered security for the Agentic components.

    OUTPUT MANDATE:
    Generate the output sequentially, adhering strictly to the five sections below.
    Format all output as valid, well-structured HTML fragments using appropriate tags (h2, h3, p, table, pre, code, etc.).
    DO NOT include <!DOCTYPE>, <html>, <body> or other full document tags - just provide the HTML content fragments.
    Start with a <h2> heading for Section 1 (not h1), and include a brief introductory paragraph explaining the core problem before Section 1.

    FUNCTIONAL REQUIREMENT DOCUMENT (FRD) INPUT:
    {document_text}

    1. Domain and Entity Identification (NIST AI RMF: MAP)

    Core Domain Definition:
    Specify the primary domain name and core focus area of the system.
    Format: "Core Domain: [Domain Name]"

    Bounded Context Definitions:
    For each bounded context (minimum 3), provide:
    - Context Name and Number (e.g., "Bounded Context 1: [Name]")
    - Focus: One-line description of the context's responsibility
    - Core Entity: List of primary entities in this context

    Entity Identification Table:
    Present in a clear tabular format:
    ```
    Entity Identification
    Context (Bounded Context Name)
    Entity (Aggregate Root)
    Key Attributes
    [Context Name]
    [Entity Name]
    [List of key attributes with types]
    ```
    
    Notes:
    - Use dimension (Dim_) and fact (Fact_) prefixes appropriately
    - Include Primary Keys (PK) and Foreign Keys (FK) in attributes
    - Follow exact format and naming conventions as shown
    - Keep descriptions concise and technical
    - Focus on structural relationships between entities

    2. Data Model Design: Database Tables and Linkages

    Table Schema: Full DDL Snippets
    For each table, provide complete CREATE TABLE statements with:
    - Table name with appropriate prefix (Dim_/Fact_)
    - All columns with precise PostgreSQL data types
    - Primary and Foreign key constraints
    - Required constraints (NOT NULL, UNIQUE)
    - Audit columns (created_by, created_at, etc.)
    - Comments explaining key constraints
    
    Present each CREATE TABLE statement in its own section, separated by a descriptive header:

    <h4>[Table Name] Definition</h4>
    <pre><code class="language-sql">
    CREATE TABLE [TableName] (
        [column_name] [data_type] [constraints],
        ...
        CONSTRAINT [constraint_name] [constraint_definition]
    );
    </code></pre>

    <p>[Brief description of table purpose and key relationships]</p>

    Relationship Diagram:
    List all relationships in format:
    [Table1] (1) ⟷ (M) [Table2] ([relationship description])

    3. Domain-Driven Design APIs (RESTful Service Contracts)

    Service List:
    List all core services with brief descriptions:
    - [ServiceName]: Brief description of service responsibility

    API Endpoints: Comprehensive Application View
    Present in a table format:
    ```
    Service | Method | Endpoint | Description | Key Roles/Authorization Scope
    ```

    Required columns:
    - Service: Name of the service
    - Method: HTTP method (GET, POST, PATCH, etc.)
    - Endpoint: Full URL path with version (/api/v1/...)
    - Description: Brief description of endpoint purpose
    - Key Roles/Authorization Scope: Required permissions

    Group endpoints by service with separators (---)

    4. Agentic AI Integration: ADK, MCP, and IPI Defense (NIST AI RMF: MANAGE)

    Agent Definition:
    Present in table format:
    ```
    Field | Detail
    Agent Name | [name]
    Agent Persona | [description]
    Primary Goal | [goal]
    ```

    Agent Development Kit (ADK) and Tool Mapping:
    List all tools in table format:
    ```
    Tool Name (Internal Function) | Mapped DDD API Endpoint | DDD Service | Purpose & Edge Case Coverage
    ```

    Model Context Protocol (MCP) Design:
    Detail the following components:

    1. Input Validation and Sanitization Layer (IPI Defense)
    - Protocol Step
    - GCP Service
    - Mechanism
    - Protocol Action

    2. Context Retrieval and Grounding Layer
    - Protocol Step
    - GCP Service
    - Mechanism

    3. Tool Orchestration and Security Wrapper
    - Protocol Step
    - GCP Service
    - Mechanism (detailed bullet points)

    4. Safety and Alignment Layer
    - Protocol Step
    - GCP Service
    - Function

    5. Frontend Pages and Agentic UX

    Page List:
    List the three most critical views with their primary purposes.

    Agent Integration Flow:
    Document the integration in table format:
    ```
    UI Element | Trigger Action | Agent Flow Step
    ```

    Security and Edge Case UX Design:
    1. AI Action Confirmation Modal (NIST AI RMF: MANAGE - AI-4.5)
    - Content requirements
    - Edge case handling
    - User feedback mechanisms

    2. Intervention and Reporting Chat UX (Principal Focus)
    Present in table format:
    ```
    UX Feature | Purpose | Edge Case Handling
    ```

    3. Policy and Grounding Transparency UX
    Detail the following:
    - Transparency mechanisms
    - Source documentation
    - Error handling
    - User feedback

    FINAL MANDATE: Review your entire output. Format everything with proper HTML tags. Ensure every instruction has been fully addressed, and that the five sections flow logically from requirements to implementation.
    """
    return prompt

def detect_file_type(file_path):
    """Detect if a file is PDF or DOCX based on its content or extension."""
    # Try to determine by content type
    content_type, _ = mimetypes.guess_type(file_path)

    if content_type == 'application/pdf':
        return 'PDF'
    elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        return 'DOCX'

    # If content type detection fails, try by extension
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == '.pdf':
        return 'PDF'
    elif ext == '.docx':
        return 'DOCX'

    return 'UNKNOWN'

# def extract_text_from_s3_file(bucket, key):
#     """Extract text from a file stored in a private S3 bucket."""
#     import hashlib
#     try:
#         # Get the file directly from S3 using the boto3 client
#         print(f"Getting file {key} from private bucket {bucket}")
#         response = s3_client.get_object(Bucket=bucket, Key=key)

#         # Read the raw file bytes
#         file_bytes = response['Body'].read()

#         # Determine file type based on key
#         file_type = detect_file_type(key)

#         text = None
#         if file_type == 'PDF':
#             # Process PDF
#             print("Processing as PDF")
#             pdf_file = io.BytesIO(file_bytes)
#             text = ""
#             with pdfplumber.open(pdf_file) as pdf:
#                 print(f"PDF has {len(pdf.pages)} pages")
#                 for page in pdf.pages:
#                     page_text = page.extract_text()
#                     if page_text:
#                         text += page_text + "\n"
#         elif file_type == 'DOCX':
#             # Process DOCX
#             print("Processing as DOCX")
#             docx_file = io.BytesIO(file_bytes)
#             doc = docx.Document(docx_file)
#             print(f"DOCX has {len(doc.paragraphs)} paragraphs")
#             text = ""
#             for para in doc.paragraphs:
#                 if para.text:
#                     text += para.text + "\n"
#         else:
#             print(f"Unsupported file type: {file_type}")
#             return {"file": key, "hash": None, "text": None}

#         # Compute hash of the extracted text for content-based deduplication
#         if text is not None:
#             # Normalize text (strip leading/trailing whitespace, collapse whitespace)
#             normalized_text = ' '.join(text.split())
#             text_hash = hashlib.sha256(normalized_text.encode('utf-8')).hexdigest()
#         else:
#             text_hash = None

#         return {"file": key, "hash": text_hash, "text": text}
#     except ClientError as e:
#         if e.response['Error']['Code'] == 'AccessDenied':
#             print(f"Access denied to S3 object: {bucket}/{key}. Check IAM permissions.")
#         elif e.response['Error']['Code'] == 'NoSuchKey':
#             print(f"File not found: {bucket}/{key}")
#         else:
#             print(f"AWS error extracting text from S3 file: {e}")
#         return {"file": key, "hash": None, "text": None}
#     except Exception as e:
#         print(f"Error extracting text from S3 file: {e}")
#         return {"file": key, "hash": None, "text": None}

def extract_text_from_s3_file(bucket, key):
    """Extract plain text from a PDF or DOCX file stored in a private S3 bucket."""
    try:
        print(f"📥 Getting file {key} from private bucket {bucket}")
        response = s3_client.get_object(Bucket=bucket, Key=key)
        file_bytes = response['Body'].read()

        file_type = detect_file_type(key)
        text = ""

        if file_type == 'PDF':
            print("📘 Processing as PDF")
            pdf_file = io.BytesIO(file_bytes)
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

        elif file_type == 'DOCX':
            print("📗 Processing as DOCX")
            docx_file = io.BytesIO(file_bytes)
            doc = docx.Document(docx_file)
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"

        else:
            print(f"❌ Unsupported file type: {file_type}")
            return {"file": key, "text": None}

        # Normalize text
        normalized_text = ' '.join(text.split()) if text else ""
        return {"file": key, "text": normalized_text}

    except ClientError as e:
        print(f"AWS S3 error: {e}")
        return {"file": key, "text": None}

    except Exception as e:
        print(f"Error extracting text from S3 file: {e}")
        return {"file": key, "text": None}

def get_model_process_stages(ai_category):
    """Get the appropriate model process stages based on AI category"""
    if ai_category == "Gen AI":
        return [
            "Requirements Analysis",
            "Data Preparation",
            "Content Processing",
            "Foundation Setup",
            "Model Configuration",
            "Prompt Engineering",
            "Implementation & Deployment",
            "Performance Monitoring",
            "Governance & Compliance"
        ]
    else:  # Machine Learning
        return [
            "Data Collection",
            "Data Preprocessing",
            "Feature Engineering",
            "Model Training",
            "Validation",
            "Deployment",
            "Monitoring",
            "Feedback Loop",
            "Retraining",
            "Audit Trail"
        ]

def generate_usecase_with_llm_focused(document_text, ai_category, cloud_provider, all_mapping_records=None):
    """
    Generate a use case using OpenAI LLM based on COMPLETE document content and AI category.
    Enhanced to generate comprehensive descriptions with both business usage and technical implementation details.
    """

    try:
        print(f"Preparing enhanced prompt for OpenAI model with category: {ai_category}")
        print(f"Using complete document text of {len(document_text)} characters")

        # --- Cloud provider extraction logic ---
        # Try to extract cloud provider from document content (case-insensitive)
        cloud_providers = ["AWS", "Azure", "GCP", "Google Cloud", "Google Cloud Platform", "Amazon Web Services", "Microsoft Azure"]
        found_provider = None
        for provider in cloud_providers:
            if provider.lower() in document_text.lower():
                found_provider = provider
                break

        # Normalize found provider
        if found_provider:
            # Map aliases to canonical names
            if found_provider in ["Google Cloud", "Google Cloud Platform"]:
                found_provider = "GCP"
            elif found_provider == "Amazon Web Services":
                found_provider = "AWS"
            elif found_provider == "Microsoft Azure":
                found_provider = "Azure"
            print(f"Cloud provider found in document: {found_provider} (priority)")
            final_cloud_provider = found_provider
        else:
            print(f"No cloud provider found in document, using passed argument: {cloud_provider}")
            final_cloud_provider = cloud_provider

        # Create category-specific system prompt and methodology guidance
        if ai_category == "Machine Learning":
            system_prompt = """You are an expert AI solution architect with deep expertise in Machine Learning methodologies. Generate comprehensive use cases that combine business strategy with technical implementation. Focus on practical, implementable solutions based on the document content provided."""
            methodology_guidance = """Focus on Machine Learning methodologies: predictive modeling, classification, regression, pattern recognition, feature engineering, statistical learning, anomaly detection."""
        elif ai_category == "AI Workflow Agents":
            system_prompt = """You are an expert AI solution architect with deep expertise in workflow automation and AI agents. Generate comprehensive use cases that combine business strategy with technical implementation. Focus on structured workflows, rule-based automation, and step-by-step orchestration."""
            methodology_guidance = """Focus on AI Workflow Agent methodologies: task automation, approval workflows, form/document processing, single-shot/few-shot execution, structured pipelines."""
        elif ai_category == "Agentic AI":
            system_prompt = """You are an expert AI solution architect with deep expertise in Agentic AI methodologies. Generate comprehensive use cases that combine business strategy with technical implementation. Focus on multi-agent collaboration, autonomous workflows, and dynamic tool selection."""
            methodology_guidance = """Focus on Agentic AI methodologies: multi-agent systems, autonomous task execution, decision orchestration, planning, agent collaboration, adaptive workflows."""
        else:  # default fallback
            system_prompt = """You are an expert AI solution architect. Generate comprehensive use cases that combine business strategy with technical implementation. Ensure business-ready, standardized outputs even if category mapping is unclear."""
            methodology_guidance = """Fallback guidance: ensure clear business-technical mapping, highlight security, compliance, and practical deployment considerations."""

        # Build the prompt
        master_info = ""
        if all_mapping_records:
            import json as _json
            # Extract unique AIMethodologyType and metrics from master records
            aimethodology_types = list({rec.get('AIMethodologyType') for rec in all_mapping_records if rec.get('AIMethodologyType')})
            metrics_master = [rec.get('metrics') for rec in all_mapping_records if rec.get('metrics')]
            # Flatten and deduplicate metrics
            flat_metrics = []
            seen = set()
            for mlist in metrics_master:
                if isinstance(mlist, list):
                    for m in mlist:
                        m_tuple = tuple(sorted(m.items())) if isinstance(m, dict) else str(m)
                        if m_tuple not in seen:
                            flat_metrics.append(m)
                            seen.add(m_tuple)
            master_info = f"""
MASTER AIMethodologyType OPTIONS (select only from these):
{_json.dumps(aimethodology_types, indent=2)}

MASTER METRICS OPTIONS (select only from these, do not invent new metrics):
{_json.dumps(flat_metrics, indent=2)}
"""

        prompt = f"""
        Based on the following COMPLETE document content analysis, generate a comprehensive AI use case specifically for a {ai_category} solution.

        COMPLETE DOCUMENT CONTENT (NO TRUNCATION):
        {document_text}

        AI CATEGORY DETERMINED: {ai_category}

        CLOUD PROVIDER TO USE: {final_cloud_provider}

        METHODOLOGY GUIDANCE: {methodology_guidance}

        {master_info}

        CRITICAL INSTRUCTIONS (STRICT):
        - All field values must be grammatically correct, human-readable, and use the correct casing and style for each field.
        - For modelInput and modelOutput, return as user-friendly, title-cased, space-separated, comma-separated strings (e.g., "Lesson Plan PPT, Student List, Assessment Report"). Do not use snake_case or underscores. Use the most natural, readable form for business users.
        - For modelName and AIMethodologyType, return a human-readable, title-cased, space-separated string (e.g., "Real Estate Consultancy AI", not "RealEstateConsultancyAI").
        - For all description and summary fields, ensure the text is grammatically correct, clear, and business-appropriate.
        - Do not invent or misspell values. For every required field, you MUST always provide the most appropriate value based on the document content and context. If a value is not explicitly mentioned, you MUST infer or select the best possible value; never leave any required field empty or omit it. All required fields are mandatory and must be present and non-empty.
        - For baseModelName, if not explicitly mentioned in the document, you MUST infer and select the most appropriate model (e.g., "GPT-4", "Claude", "Gemini", "Bedrock", etc.) based on the use case and always provide a value. baseModelName must never be empty.
        - All output must be valid JSON parsable by json.loads().

        REQUIRED FIELDS (ALL MUST BE DERIVED FROM THE DOCUMENT CONTENT):

        1. businessUsage: Extract the specific business application from the document content - provide 2-3 detailed, grammatically correct sentences describing exactly what is mentioned in the document.
        2. currentBusinessUsage: Describe how the process or activity is currently being performed in the real business before AI implementation. Focus on manual steps, inefficiencies, challenges, or existing workflows that the AI solution aims to improve. Provide 2-3 clear, grammatically correct sentences that reflect the current or existing state of the business process.
        3. department: Identify the department mentioned in the document or infer from the business context described.
        4. usecaseCategory: Set to "{ai_category}".
        5. impact: Assess business impact based on what's described in the document:
         - HIGH: Strategic transformation, significant efficiency gains, competitive advantage
         - MEDIUM: Operational improvements, moderate efficiency gains
         - LOW: Minor enhancements, limited scope
        6. level: Determine AI maturity level (0-6) based on the complexity and scope described in the document:
        {level_finding_content}
        7. AIMethodologyType: Select the most appropriate value from the MASTER AIMethodologyType OPTIONS above and return it as a human-readable, title-cased, space-separated string (e.g., "Real Estate Consultancy AI").
        8. baseModelName: Identify the most appropriate base model (e.g., Gemini, Bedrock, GPT-4, Claude, etc.) mentioned in the document or suitable for the described use case. If not explicitly mentioned, you MUST infer and select the most appropriate model and always provide a value. baseModelName must never be empty.
        9. keyActivity: Extract the primary activity/function described in the document in 1-2 concise, grammatically correct sentences.
        10. modelInput: Extract all relevant input data types mentioned or implied in the document. If multiple, return as a single comma-separated string in user-friendly, title-cased, space-separated format (e.g., "Lesson Plan PPT, Student List, Assessment Report"). Use correct spelling and casing.
        11. modelOutput: Extract all relevant output data types mentioned or implied in the document. If multiple, return as a single comma-separated string in user-friendly, title-cased, space-separated format (e.g., "PDF Report, Email Notification"). Use correct spelling and casing.
        12. modelName: Create a human-readable, title-cased, space-separated name based on the use case described in the document (e.g., "Property Requirements AI", "Real Estate Consultancy AI").
        13. modelDescription: Based on the document content, provide a comprehensive 3-4 sentence description of the model's functionality, ensuring all details are derived from the document and the text is grammatically correct.
        14. modelSummary: Provide a concise 1-2 sentence summary based on the document content, ensuring all details are derived from the document and the text is grammatically correct.
        15. modelPurpose: Extract the specific goal/purpose from the document content.
        16. modelUsage: Based on the document, describe how the model would be used technically.
        17. overallRisk: Based on the document content, assess the overall risk level:
            - LOW: Basic automation, minimal sensitive data, simple business processes
            - MEDIUM: Moderate automation, some personal data, industry compliance considerations
            - HIGH: Critical decisions, extensive sensitive data, high regulatory requirements, safety-critical
            - PROHIBITED: Biometric surveillance, social scoring, behavioral manipulation
        18. platform: Extract or infer relevant technologies from the document content as a comma-separated string, and always append the cloud provider ({final_cloud_provider}) at the end (e.g., "Python, FastAPI, {final_cloud_provider}").
        19. cloudProvider: Set to the cloud provider found in the document (if present), otherwise use the passed argument ({final_cloud_provider}).
        20. priorityType: Determine the risk classification based on document content analysis:
            - "LOW-RISK AI": For basic business automation with minimal risk factors
            - "MEDIUM-RISK AI": For moderate complexity systems with some risk considerations
            - "HIGH-RISK AI": For critical systems with significant risk factors or regulatory requirements
            - "PROHIBITED AI": Only if document describes prohibited use cases (rare)
        21. sector: Extract or infer the industry sector from the document content.
        22. useFrequency: Extract or infer usage frequency from the document content. Only select one of the following master values: "Daily", "Weekly", "Monthly", or "Yearly".

        23. metrics: Select the most appropriate metrics (5-7 total). Prioritize selection from the MASTER METRICS OPTIONS. However, to ensure **real business measurement** and not just generic AI metrics, you may **infer and append 1-2 additional, highly specific metrics** that are crucial to the documented business use case, if the master list lacks necessary business relevance. For each selected metric, set a realistic threshold value (0-100) based on the document content and the identified AIMethodologyType (do not use any threshold from master data, always infer from the document and context).

        CRITICAL TIP FOR METRICS: Metrics must be highly relevant and tied to tangible business outcomes, not just model accuracy. Use the following examples to guide your specificity:

        - **GENERIC (AVOID):** "F1 Score," "Accuracy" (unless strictly necessary)
        - **REAL ESTATE (EXAMPLE):** "Property Listing Match Rate (85%)," "Lead-to-View Conversion Rate (70%)," "Time-to-Offer Reduction (75%)"
        - **HEALTHCARE CLAIMS (EXAMPLE):** "Claims Processing Straight-Through Rate (90%)," "Error Rate on Approved Claims (98%)," "Time to Adjudication Reduction (75%)"
        - **EDUCATION (EXAMPLE):** "Student Assessment Consistency (95%)," "Lesson Plan Generation Time (90%)," "Teacher Time Saved Per Week (80%)"
        - **INTERVIEW AGENT (EXAMPLE):** "Candidate Quality Rating Consistency (85%)," "Screening Time Per Candidate Reduction (70%)," "Hiring Manager Satisfaction Score (80%)"

        Each metric must follow this JSON structure:
        [
            {{
                "metricName": "",
                "metricDescription": "",
                "threshold": 0
            }}
        ]

        24. status: Set to "Not yet started".
        25. searchAttributesAsJson: Combine key attributes derived from document as comma-separated string.
        26. questions: Extract ALL explicit questions that are specifically written in the document. Look for:
            - Questions ending with "?"
            - Interrogative statements
            - Requests for information or clarification
            - Any queries posed in the document
            Do NOT generate new questions - only extract the actual questions that appear in the document text.
            Return as an array of strings containing the exact questions found in the document.
        27. isProposalGenerated: Set to false.

        CRITICAL FORMATTING REQUIREMENTS:
        - Return valid JSON that can be parsed by json.loads().
        - All field values must be grammatically correct, human-readable, and use the correct casing and style for each field (e.g., Real Estate Consultancy AI, property_requirements_form, PDF_report).
        - Do not invent, misspell, or change casing of any value. Use master options and document content exactly as provided.
        - Include specific {ai_category} tools, technologies, and implementation details in all relevant fields.
        - Show clear business value progression alongside technical development.
        - ALL content must be derived from the actual document content provided above.
        - Focus on {ai_category}-specific approaches and methodologies.
        - Ensure accuracy to the document content - no generic responses.

        Generate the complete JSON response now:
        """

        print("Calling OpenAI API for enhanced usecase generation with complete document...")
        # Use higher max_tokens for complete document processing with enhanced descriptions
        response = call_llm(prompt, system_prompt, max_tokens=7000)
        print("Processing enhanced OpenAI response...")
        generated_text = response

        print("Extracting and validating JSON from response...")
        # Extract JSON with multiple fallback strategies
        try:
            # Try to parse as JSON directly first
            usecase_data = json.loads(generated_text)
        except json.JSONDecodeError:
            # If direct parsing fails, try to extract JSON from response
            json_start = generated_text.find('{')
            json_end = generated_text.rfind('}') + 1
            if json_start != -1 and json_end > json_start:
                json_text = generated_text[json_start:json_end]
                try:
                    usecase_data = json.loads(json_text)
                except json.JSONDecodeError:
                    # Try cleaning up common JSON issues
                    json_text = json_text.replace('\\"', '"').replace('\\n', '\n')
                    # Remove any trailing commas
                    json_text = re.sub(r',(\s*[}\]])', r'\1', json_text)
                    try:
                        usecase_data = json.loads(json_text)
                    except json.JSONDecodeError as e:
                        print(f"Failed to parse JSON after cleanup: {e}")
                        print(f"JSON text: {json_text[:500]}...")
                        raise Exception(f"Invalid JSON format in response: {e}")
            else:
                raise Exception("No valid JSON found in response")

        except Exception as e:
            print(f"JSON parsing error: {e}")
            raise Exception(f"Error parsing LLM response: {e}")

        # Ensure the category matches what was determined
        usecase_data['usecaseCategory'] = ai_category
        # Ensure cloudProvider field is set correctly
        usecase_data['cloudProvider'] = final_cloud_provider

        # Process and enhance the data
        usecase_data = enhance_usecase_data(usecase_data, ai_category)

        return usecase_data

    except Exception as e:
        print(f"Error in enhanced usecase generation: {e}")
        return {"error": str(e)}

def enhance_usecase_data(usecase_data, ai_category):
    """Enhance and validate usecase data similar to usecase-creation lambda"""
    try:
        # Convert floats to Decimal for DynamoDB
        usecase_data = convert_floats_to_decimal_for_dynamodb(usecase_data)

        # Ensure usecase_data is a dict
        if not isinstance(usecase_data, dict):
            print("Warning: usecase_data is not a dict, returning empty structure")
            usecase_data = {}

        # Ensure usecaseCategory matches determined AI category
        usecase_data["usecaseCategory"] = ai_category

        # Add searchAttributesAsJson field safely
        usecase_data["searchAttributesAsJson"] = ",".join([
            str(usecase_data.get("AIMethodologyType", "")),
            str(usecase_data.get("modelName", "")),
            str(usecase_data.get("sector", "")),
            str(usecase_data.get("platform", ""))
        ])

        # Handle metrics safely
        raw_metrics = usecase_data.get("metrics", [])
        metrics = []

        if isinstance(raw_metrics, list) and raw_metrics:
            # Preserve OpenAI metrics, ensure threshold exists
            for m in raw_metrics:
                if isinstance(m, dict):
                    if "threshold" not in m:
                        m["threshold"] = 80  # fallback default
                    metrics.append(m)
        else:
            # Set default metrics based on AI category
            if ai_category == "AI Workflow Agents":
                metrics = [
                    {"metricName": "Response Relevance", "metricDescription": "Measure of how well the agent follows the defined workflow steps", "threshold": 85},
                    {"metricName": "Workflow Accuracy", "metricDescription": "Percentage of tasks executed correctly within the defined workflow", "threshold": 90},
                    {"metricName": "User Satisfaction", "metricDescription": "Rating of agent interaction quality and usefulness", "threshold": 90},
                    {"metricName": "Response Time", "metricDescription": "Time taken for the agent to complete a workflow step", "threshold": 5}
                ]
            elif ai_category == "Machine Learning":
                metrics = [
                    {"metricName": "Accuracy", "metricDescription": "The proportion of correctly classified instances out of the total instances", "threshold": 90},
                    {"metricName": "Precision", "metricDescription": "The ratio of correctly predicted positive instances to all predicted positives", "threshold": 85},
                    {"metricName": "Recall", "metricDescription": "The ratio of correctly predicted positive instances to actual positives", "threshold": 85},
                    {"metricName": "F1-Score", "metricDescription": "The harmonic mean of precision and recall", "threshold": 85},
                    {"metricName": "ROC-AUC", "metricDescription": "Area under the Receiver Operating Characteristic curve", "threshold": 90}
                ]
            elif ai_category == "Agentic AI":
                metrics = [
                    {"metricName": "Task Completion Rate", "metricDescription": "Percentage of tasks successfully executed by agents", "threshold": 90},
                    {"metricName": "Decision Accuracy", "metricDescription": "Accuracy of autonomous decisions or recommendations", "threshold": 85},
                    {"metricName": "System Reliability", "metricDescription": "Measure of uptime and fault tolerance of the AI system", "threshold": 95},
                    {"metricName": "Collaboration Efficiency", "metricDescription": "Effectiveness of multi-agent coordination and workflow execution", "threshold": 85},
                    {"metricName": "Adaptability", "metricDescription": "How effectively agents adapt to changing goals or environments", "threshold": 80}
                ]
            else:
                # Safe fallback if category is missing or new
                metrics = [
                    {"metricName": "Generic Success Rate", "metricDescription": "Fallback metric when category is unclear", "threshold": 80}
                ]

        usecase_data["metrics"] = metrics

        # Ensure required fields are strings
        string_fields = ["platform", "businessUsage", "modelDescription"]
        for field in string_fields:
            val = usecase_data.get(field, "")
            if isinstance(val, list):
                usecase_data[field] = ", ".join(str(item) for item in val)
            elif not isinstance(val, str):
                usecase_data[field] = str(val)

        # Set required status
        usecase_data["status"] = "Not yet started"

        # Remove unwanted fields
        unwanted_fields = ["processingStatus", "documentType", "userId", "modelProcess"]
        for field in unwanted_fields:
            if field in usecase_data:
                del usecase_data[field]

        print("Successfully enhanced usecase data")
        return usecase_data

    except Exception as e:
        print(f"Error enhancing usecase data: {e}")
        return {"error": f"Error enhancing data: {e}"}

def generate_sequential_id(dynamodb_resource, table_name, prefix):
    """
    Generate a sequential ID by finding and incrementing the highest existing ID with the given prefix.
    Important for maintaining consistent ID patterns across the application.
    """
    print(f"Generating sequential ID with prefix '{prefix}' for table: {table_name}")

    # Get the table
    table = dynamodb_resource.Table(table_name)

    try:
        # Scan the table to find the highest existing ID with the given prefix
        print(f"Scanning table for existing IDs with prefix: {prefix}")
        response = table.scan(
            ProjectionExpression="id",
            FilterExpression="begins_with(id, :prefix)",
            ExpressionAttributeValues={':prefix': prefix}
        )

        items = response['Items']
        print(f"Found {len(items)} items with prefix {prefix}")

        # Handle pagination if needed
        while 'LastEvaluatedKey' in response:
            print("Pagination detected, fetching more items")
            response = table.scan(
                ProjectionExpression="id",
                FilterExpression="begins_with(id, :prefix)",
                ExpressionAttributeValues={':prefix': prefix},
                ExclusiveStartKey=response['LastEvaluatedKey']
            )
            new_items = response['Items']
            print(f"Found {len(new_items)} additional items")
            items.extend(new_items)

        # Find the highest number
        max_num = 0
        prefix_len = len(prefix)
        print("Finding highest ID number")

        for item in items:
            try:
                id_str = item.get('id', '')
                if id_str.startswith(prefix):
                    num_part = id_str[prefix_len:]  # Extract the numeric part
                    if num_part.isdigit():
                        current_num = int(num_part)
                        if current_num > max_num:
                            max_num = current_num
                            print(f"New highest number found: {max_num}")
            except (ValueError, IndexError) as e:
                print(f"Error parsing ID {item.get('id', '')}: {str(e)}")
                continue

        # Generate new ID with incremented number - 3 digits
        new_num = max_num + 1
        new_id = f"{prefix}{new_num:03d}"  # Format with 3 leading zeros
        print(f"Generated new sequential ID: {new_id}")

        return new_id

    except Exception as e:
        print(f"Error generating sequential ID: {str(e)}")
        # Fallback to a default pattern with timestamp
        timestamp = int(datetime.now().timestamp())
        fallback_id = f"{prefix}{timestamp}"
        print(f"Using fallback ID with timestamp: {fallback_id}")
        return fallback_id

def is_duplicate_document_hash(table, document_hash):
    """Check if the given document_hash already exists in the usecase table."""
    if not document_hash:
        return False
    try:
        # Scan for the hash value in the table
        response = table.scan(
            FilterExpression=Attr('documentHash').eq(document_hash),
            ProjectionExpression="documentHash"
        )
        items = response.get('Items', [])
        if items:
            return True
        return False
    except Exception as e:
        print(f"Error checking duplicate document hash: {e}")
        return False
    
# def process_s3_event_record(record, dynamodb_resource):
#     """Process a single S3 event record from a private S3 bucket"""
#     try:
#         # Extract bucket and key information
#         bucket = record['s3']['bucket']['name']
#         key = urllib.parse.unquote_plus(record['s3']['object']['key'])

#         # Get AWS region dynamically from the record (or fallback to boto3 default region)
#         region = record.get("awsRegion") or Session().region_name

#         # Build full S3 URL
#         full_url = f"https://{bucket}.s3.{region}.amazonaws.com/{key}"
#         print(f"Processing S3 object from private bucket: {full_url}")

#         # Extract COMPLETE text and hash directly from S3 (no truncation)
#         doc_info = extract_text_from_s3_file(bucket, key)
#         print(f"hash={doc_info.get('hash')}, text_length={len(doc_info.get('text') or '')}")
#         if not doc_info or not doc_info.get('text'):
#             error_msg = "Failed to extract text from document or document is empty"
#             print(error_msg)
#             return {
#                 'file': key,
#                 'error': error_msg
#             }

#         document_text = doc_info['text']
#         document_hash = doc_info['hash']

#         # Extract app_name and stage_name from event
#         event = {'Records': [record]}
#         app_name, stage_name, _ = extract_app_and_stage_from_event(event)

#         print(f"Extracted app_name: {app_name}, stage_name: {stage_name}")

#         # Get table name and framework ID
#         usecase_table_name, framework_id = get_usecase_assessment_table_and_framework_id(event, dynamodb_resource)

#         # Check for duplicate document hash in the usecase table
#         usecase_table = dynamodb_resource.Table(usecase_table_name)
#         is_duplicate = is_duplicate_document_hash(usecase_table, document_hash)
#         if is_duplicate:
#             print(f"Duplicate document content detected for hash: {document_hash}")
#             return {
#                 'file': key,
#                 'error': f'Duplicate document content detected for hash: {document_hash}'
#             }
#         else:
#             print(f"No duplicate document content found for hash: {document_hash}")

#         # Methodology Metrics Mapping Table
#         methodology_metrics_mapping_table = f"{stage_name}-{app_name}-methodologyMetricsMapping"
#         all_mapping_records = fetch_all_dynamodb_records(methodology_metrics_mapping_table, dynamodb)


#         document_summary = summarize_document_with_llm(document_text)["documentSummary"]
#         print(f"Document summary: {document_summary}")

#         # Step 1: Determine AI category using COMPLETE document
#         print("\nStep 1: Determining AI category using complete document...")
#         try:
#             result = determine_ai_category_and_approach_with_llm(document_text)
#             ai_category = result["category"]
#             ai_approach = result["approach"]
#             raw_response = result["raw_response"]
#             print(f"Determined AI category: {ai_category}, AI approach: {ai_approach}, Raw response: {raw_response}")
#         except Exception as e:
#             print(f"Error in AI category determination: {e}")
#             # Fallback values same as method defaults
#             ai_category = "AI Workflow Agents"
#             ai_approach = "Custom AI Stack"
#             raw_response = "{Fallback due to error}"
#             print(f"Falling back to default AI category: {ai_category}, AI approach: {ai_approach}")

#         # Step 2: Generate enhanced use case using COMPLETE document
#         print(f"\nStep 2: Generating enhanced use case for {ai_category} using complete document...")
#         cloud_provider = "GCP"
#         usecase_data = generate_usecase_with_llm_focused(document_text, ai_category, cloud_provider, all_mapping_records)

#         if "error" in usecase_data:
#             print(f"Error in use case generation: {usecase_data['error']}")
#             return {'file': key, 'error': usecase_data['error']}

#         # Save use case to DynamoDB
#         usecase_id = generate_sequential_id(dynamodb_resource, usecase_table.name, "AI-UC-AST-")
#         timestamp = datetime.utcnow().isoformat()

#         usecase_item = {
#             "id": usecase_id,
#             "createdAt": timestamp,
#             "updatedAt": timestamp,
#             #"documentSource": f"s3://{bucket}/{key}",
#             "sourceDocURL": full_url,
#             "category": "AI Inventory",
#             "processingStatus": "completed",
#             "riskframeworkid": framework_id,
#             "aiApproach": ai_approach,
#             "aiCategory": ai_category,
#             "documentHash": document_hash,
#             "documentSummary": document_summary,
#             "aiCloudProvider": usecase_data.get("cloudProvider", cloud_provider),
#             **{k: v for k, v in usecase_data.items() if k not in ["documentType", "userId"]}
#         }

#         print(f"Saving enhanced use case with ID {usecase_id}...")
#         usecase_table.put_item(Item=usecase_item)
#         print("Successfully saved enhanced use case")

#         return {
#             'file': key,
#             'result': {
#                 'usecase_id': usecase_id,
#                 'ai_category': ai_category,
#                 'status': 'success'
#             }
#         }

#     except ClientError as e:
#         # Handle specific AWS errors like access denied
#         if e.response['Error']['Code'] == 'AccessDenied':
#             print(f"Access denied to S3 object: {bucket}/{key}") # type: ignore
#             return {
#                 'file': key, # type: ignore
#                 'error': f"Access denied to S3 object: {bucket}/{key}. Check IAM permissions." # type: ignore
#             }
#         else:
#             print(f"AWS error processing S3 event record: {e}")
#             return {
#                 'file': key, # type: ignore
#                 'error': str(e)
#             }
#     except Exception as e:
#         print(f"Error processing S3 event record: {e}")
#         return {
#             'file': key if 'key' in locals() else 'unknown', # type: ignore
#             'error': str(e)
#         }

def process_usecase_request(hash_value, s3url, cloud_id, dynamodb_resource, riskframeworkid=None):
    """
    Process a new use case request using precomputed hash and S3 file URL.

    Steps:
    1. Parse S3 bucket and object key from URL
    2. Derive app_name, stage_name, tenant_id_suffix from bucket name
    3. Extract text from S3 file
    4. Determine table names and framework ID
    5. Fetch methodology metrics mapping records
    6. Summarize document
    7. Determine AI category and approach using LLM
    8. Generate enhanced use case using LLM
    9. Save the use case to DynamoDB
    """

    bucket = "unknown"
    key = "unknown"

    try:
        # Step 1: Parse bucket and key
        parsed = urllib.parse.urlparse(s3url)
        bucket = parsed.netloc.split(".s3")[0]
        key = parsed.path.lstrip("/")
        print(f"🔹 Processing file from S3 URL: {s3url}")
        print(f"✅ Extracted bucket: {bucket}, key: {key}")
        print(f"🔹 Provided hash: {hash_value}")

        # Step 2: Extract app_name, stage_name, tenant_id_suffix
        app_name, stage_name, tenant_id_suffix = extract_app_stage_tenant_from_bucket(bucket)
        print(f"✅ Derived app_name={app_name}, stage_name={stage_name}, tenant_id_suffix={tenant_id_suffix}")
        
        # tenant_id_suffix is not used further in this function

        # Step 3: Extract document text from S3
        doc_info = extract_text_from_s3_file(bucket, key)
        document_text = doc_info.get("text")
        if not document_text:
            return {"s3url": s3url, "error": "Failed to extract text or empty document"}
        print(f"📄 Extracted text length: {len(document_text)} characters")

        # Step 4: Generate design document
        print("Generating design document prompt...")
        design_doc_prompt = generate_design_document_prompt(document_text)

        # Step 5: Call OpenAI to generate design document
        print("Calling OpenAI to generate design document...")
        system_prompt = "You are a Solution Architect and Agentic AI Specialist generating a comprehensive technical design document."
        design_doc_content = call_llm(design_doc_prompt, system_prompt, use_json_format=False, max_tokens=7000)
        print("Successfully generated design document")

        # Step 6: Determine usecase table name and framework ID
        usecase_table_name, framework_id = get_usecase_assessment_table_and_framework_id(
            app_name, stage_name, cloud_id, dynamodb_resource
        )
        print(f"Resolved usecase table: {usecase_table_name}, framework_id: {framework_id}")

        # Step 7: Fetch methodology metrics mapping records
        methodology_metrics_mapping_table = f"{stage_name}-{app_name}-methodologyMetricsMapping"
        all_mapping_records = fetch_all_dynamodb_records(methodology_metrics_mapping_table, dynamodb_resource)

        # Step 8: Summarize document
        document_summary = summarize_document_with_llm(document_text)["documentSummary"]
        print(f"Document summary: {document_summary}")

        # Step 9: Determine AI category and approach
        print("\nStep 9: Determining AI category using complete document...")
        try:
            result = determine_ai_category_and_approach_with_llm(document_text)
            ai_category = result["category"]
            ai_approach = result["approach"]
            raw_response = result["raw_response"]
        except Exception as e:
            print(f"Error determining AI category: {e}")
            ai_category = "AI Workflow Agents"
            ai_approach = "Custom AI Stack"
            raw_response = "{Fallback due to error}"

        print(f"AI category: {ai_category}, AI approach: {ai_approach}, Raw response: {raw_response}")

        # Step 10: Generate enhanced use case
        cloud_provider = "GCP"
        usecase_data = generate_usecase_with_llm_focused(document_text, ai_category, cloud_provider, all_mapping_records)
        if "error" in usecase_data:
            return {'file': key, 'error': usecase_data['error']}

        # Step 11: Save use case to DynamoDB
        usecase_table = dynamodb_resource.Table(usecase_table_name)
        usecase_id = generate_sequential_id(dynamodb_resource, usecase_table.name, "AI-UC-AST-")
        timestamp = datetime.utcnow().isoformat()
        print(f"Generated usecase ID: {usecase_id}")
        print(f"Risk Framework ID to use: {riskframeworkid if riskframeworkid else framework_id}")  
        usecase_item = {
            "id": usecase_id,
            "createdAt": timestamp,
            "updatedAt": timestamp,
            "sourceDocURL": s3url,
            "category": "AI Inventory",
            "processingStatus": "completed",
            "riskframeworkid": riskframeworkid if riskframeworkid else framework_id,
            "aiApproach": ai_approach,
            "aiCategory": ai_category,
            "documentHash": hash_value,
            "documentSummary": document_summary,
            "designDocument": design_doc_content,
            "aiCloudProvider": usecase_data.get("cloudProvider", cloud_provider),
            **{k: v for k, v in usecase_data.items() if k not in ["documentType", "userId"]}
        }
        print(f"Saving enhanced use case with ID {usecase_id}...")
        usecase_table.put_item(Item=usecase_item)
        print("Successfully saved enhanced use case")

        return {
            'file': key,
            'result': {'usecase_id': usecase_id, 'ai_category': ai_category, 'status': 'success'}
        }

    except ClientError as e:
        error_code = e.response['Error']['Code']
        if error_code == 'AccessDenied':
            print(f"Access denied to S3 object: {bucket}/{key}")
            return {'file': key, 'error': f"Access denied to S3 object: {bucket}/{key}. Check IAM permissions."}
        else:
            print(f"AWS ClientError: {e}")
            return {'file': key, 'error': str(e)}

    except Exception as e:
        print(f"Unexpected error: {e}")
        return {'file': key if 'key' in locals() else 'unknown', 'error': str(e)}

# Lambda handler function
# def lambda_handler(event, context):
#     """
#     AWS Lambda handler function to process documents from S3 events.
#     Processes uploaded files, extracts text, generates use cases via LLM, and stores in DynamoDB.
#     """
#     print(f"Received event: {json.dumps(event)}")

#     try:
#         # Check if this is an S3 event
#         if 'Records' in event and len(event['Records']) > 0 and event['Records'][0].get('eventSource') == 'aws:s3':
#             print("Processing S3 event with multiple records")
#             results = []
#             for record in event['Records']:
#                 record_result = process_s3_event_record(record, dynamodb)
#                 results.append(record_result)
#             return {
#                 'statusCode': 200,
#                 'body': json.dumps({
#                     'message': f'Successfully processed {len(results)} files',
#                     'results': results
#                 })
#             }
#         else:
#             return {
#                 'statusCode': 400,
#                 'body': json.dumps({
#                     'message': 'Invalid event format. Expected S3 event trigger.'
#                 })
#             }
#     except Exception as e:
#         print(f"Unexpected error in lambda_handler: {e}")
#         return {
#             'statusCode': 500,
#             'body': json.dumps({'error': str(e)})
#         }


def lambda_handler(event, _context):
    _ = _context  # mark as intentionally unused
    
    """
    AWS Lambda handler function (POST API version).
    Receives 'hash', 's3url', and 'cloudId' from API Gateway request,
    validates input, and triggers use case creation logic.
    """
    import json

    print(f"Received event: {json.dumps(event)}")

    try:
        # Parse and validate body
        body = json.loads(event.get("body", "{}"))
        hash_value = body.get("hash")
        s3url = body.get("s3url")
        cloud_id = body.get("cloudId")
        riskframeworkid = body.get("riskframeworkid")  # optional

        if not hash_value or not s3url or not cloud_id:
            return {
                'statusCode': 400,
                'body': json.dumps({
                    'message': 'Fields "hash", "s3url", and "cloudId" are required in the request body.'
                })
            }

        print(f"✅ Received hash: {hash_value}")
        print(f"✅ Received S3 URL: {s3url}")
        print(f"✅ Received cloudId: {cloud_id}")
        print(f"✅ Received riskframeworkid: {riskframeworkid}")

        # 🔹 Call internal processing logic
        result = process_usecase_request(hash_value, s3url, cloud_id, dynamodb, riskframeworkid)  # cloudId passed

        return {
            'statusCode': 200,
            'body': json.dumps({
                'message': 'Usecase processed successfully',
                'result': result
            })
        }

    except Exception as e:
        print(f"Error in lambda_handler: {e}")
        return {
            'statusCode': 500,
            'body': json.dumps({'error': str(e)})
        }
